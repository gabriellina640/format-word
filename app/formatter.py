from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.config import FormatSettings


SUPPORTED_INPUT_EXTENSIONS = {".docx", ".pdf"}
HEADER_IMAGE_WIDTH_CM = 15.5
HEADER_IMAGE_HEIGHT_CM = 2.1
FOOTER_IMAGE_WIDTH_CM = 15.5
FOOTER_IMAGE_HEIGHT_CM = 1.7


class FormatterError(Exception):
    """Raised when a document cannot be formatted safely."""


@dataclass(slots=True)
class FormatResult:
    output_path: Path
    paragraphs: int


def format_document(input_path: Path, output_dir: Path, settings: FormatSettings) -> FormatResult:
    input_path = input_path.expanduser().resolve()
    output_dir = output_dir.expanduser().resolve()

    validate_input(input_path, settings.max_input_mb)
    _prepare_output_dir(output_dir)

    paragraphs = read_document_paragraphs(input_path)
    if not paragraphs:
        raise FormatterError("Não foi possível extrair texto do arquivo informado.")

    return format_paragraphs(paragraphs, output_dir, input_path.stem, settings)


def format_paragraphs(
    paragraphs: list[str],
    output_dir: Path,
    input_stem: str,
    settings: FormatSettings,
) -> FormatResult:
    output_dir = output_dir.expanduser().resolve()
    _prepare_output_dir(output_dir)

    cleaned_paragraphs = [paragraph.strip() for paragraph in paragraphs if paragraph.strip()]
    if not cleaned_paragraphs:
        raise FormatterError("Não há conteúdo para exportar.")

    document = _create_output_document(settings)
    _configure_sections(document, settings)
    _apply_header_footer(document, settings)
    _write_paragraphs(document, cleaned_paragraphs, settings)

    output_path = _unique_output_path(output_dir, input_stem, settings.output_suffix)
    document.save(output_path)
    return FormatResult(output_path=output_path, paragraphs=len(cleaned_paragraphs))


def _create_output_document(settings: FormatSettings) -> Document:
    template_path = Path(settings.template_path) if settings.template_path else None
    if template_path and template_path.is_file():
        document = Document(str(template_path))
        _clear_document_body(document)
        return document
    return Document()


def _uses_template(settings: FormatSettings) -> bool:
    return bool(settings.template_path and Path(settings.template_path).is_file())


def _clear_document_body(document: Document) -> None:
    body = document._body._element
    section_properties = body.sectPr
    for child in list(body):
        if child is section_properties:
            continue
        body.remove(child)


def validate_input(input_path: Path, max_input_mb: int) -> None:
    if input_path.suffix.lower() not in SUPPORTED_INPUT_EXTENSIONS:
        raise FormatterError("Selecione um arquivo .docx ou .pdf.")
    if not input_path.is_file():
        raise FormatterError("Arquivo não encontrado.")
    if input_path.stat().st_size > max_input_mb * 1024 * 1024:
        raise FormatterError(f"O arquivo deve ter no máximo {max_input_mb} MB.")


def _prepare_output_dir(output_dir: Path) -> None:
    if output_dir.exists() and not output_dir.is_dir():
        raise FormatterError("A pasta de saída selecionada não é um diretório.")
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise FormatterError(f"Não foi possível acessar a pasta de saída: {exc}") from exc


def read_document_paragraphs(input_path: Path) -> list[str]:
    input_path = input_path.expanduser().resolve()
    if input_path.suffix.lower() == ".docx":
        return _read_docx_paragraphs(input_path)
    if input_path.suffix.lower() == ".pdf":
        return _read_pdf_paragraphs(input_path)
    return []


def _read_docx_paragraphs(input_path: Path) -> list[str]:
    document = Document(input_path)
    content = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_inline_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                content.append(row_text)
    return content


def _read_pdf_paragraphs(input_path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FormatterError("Instale a dependência pypdf para importar arquivos PDF.") from exc

    reader = PdfReader(str(input_path))
    if reader.is_encrypted:
        raise FormatterError("PDF protegido por senha não é suportado.")

    paragraphs: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for block in text.split("\n\n"):
            cleaned = " ".join(line.strip() for line in block.splitlines() if line.strip())
            if cleaned:
                paragraphs.append(cleaned)
    return paragraphs


def _normalize_inline_text(text: str) -> str:
    return " ".join(part.strip() for part in text.splitlines() if part.strip())


def _configure_sections(document: Document, settings: FormatSettings) -> None:
    section = document.sections[0]
    section.top_margin = Cm(settings.margin_top_cm)
    section.bottom_margin = Cm(settings.margin_bottom_cm)
    section.left_margin = Cm(settings.margin_left_cm)
    section.right_margin = Cm(settings.margin_right_cm)
    section.header_distance = Cm(max(0.3, 0.8 + settings.header_offset_y_cm))
    section.footer_distance = Cm(max(0.3, 0.7 - settings.footer_offset_y_cm))

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = settings.font_name
    normal_style.font.size = Pt(settings.font_size)
    normal_style.paragraph_format.line_spacing = settings.line_spacing
    normal_style.paragraph_format.space_after = Pt(settings.paragraph_spacing_after)
    normal_style.paragraph_format.first_line_indent = Cm(settings.first_line_indent_cm)
    normal_style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if settings.justify_text else WD_ALIGN_PARAGRAPH.LEFT
    )


def _apply_header_footer(document: Document, settings: FormatSettings) -> None:
    if settings.template_path:
        return

    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    if settings.include_header and settings.header_image_path:
        _add_section_image(
            section.header.paragraphs[0],
            Path(settings.header_image_path),
            available_width,
            settings.header_offset_x_cm,
            HEADER_IMAGE_WIDTH_CM,
            HEADER_IMAGE_HEIGHT_CM,
        )
    if settings.include_footer and settings.footer_image_path:
        _add_section_image(
            section.footer.paragraphs[0],
            Path(settings.footer_image_path),
            available_width,
            settings.footer_offset_x_cm,
            FOOTER_IMAGE_WIDTH_CM,
            FOOTER_IMAGE_HEIGHT_CM,
        )


def _add_section_image(
    paragraph,
    image_path: Path,
    available_width,
    offset_x_cm: float,
    width_cm: float,
    height_cm: float,
) -> None:
    if not image_path.is_file():
        return
    available_width_cm = available_width / Cm(1)
    base_indent_cm = max(0.0, (available_width_cm - width_cm) / 2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(max(0.0, base_indent_cm + offset_x_cm))
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))


def _write_paragraphs(document: Document, paragraphs: list[str], settings: FormatSettings) -> None:
    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if settings.justify_text else WD_ALIGN_PARAGRAPH.LEFT
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = settings.line_spacing
        paragraph_format.space_after = Pt(settings.paragraph_spacing_after)
        paragraph_format.first_line_indent = Cm(settings.first_line_indent_cm)

        run = paragraph.add_run(text)
        run.font.name = settings.font_name
        run.font.size = Pt(settings.font_size)


def _unique_output_path(output_dir: Path, input_stem: str, suffix: str) -> Path:
    safe_suffix = _sanitize_suffix(suffix)
    base = output_dir / f"{input_stem}{safe_suffix}.docx"
    if not base.exists():
        return base

    counter = 2
    while True:
        candidate = output_dir / f"{input_stem}{safe_suffix}_{counter}.docx"
        if not candidate.exists():
            return candidate
        counter += 1


def _sanitize_suffix(suffix: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", suffix.strip())
    cleaned = cleaned.strip("_")
    return f"_{cleaned}" if cleaned else "_formatado"
